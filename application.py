from flask import Flask, render_template, request, send_file, redirect, url_for, flash, session
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from docxtpl import DocxTemplate
from docx import Document
import datetime
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx2pdf import convert
import os
import mammoth
from pymongo import MongoClient
from gridfs import GridFS
import logging
import pythoncom

application = Flask(__name__, static_url_path='/static')
app = application
#app.secret_key = 'your_secret_key'

# Configure logging
logging.basicConfig(level=logging.DEBUG)

# MongoDB setup
client = MongoClient('mongodb+srv://beingchirag0051:Chirag5151@cluster0.e6xh9gf.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0')
db = client['beingchirag0051']
fs = GridFS(db)
pdfs_metadata = db['pdfs_metadata']

# Flask-Login setup
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# User class for Flask-Login
class User(UserMixin):
    def __init__(self, id):
        self.id = id

# In-memory user storage for simplicity
users = {
    'admin': {'password': 'admin_pass'},
    'user1': {'password': 'user1_pass'},
    'user2': {'password': 'user2_pass'}
}

@login_manager.user_loader
def load_user(user_id):
    return User(user_id) if user_id in users else None

@app.route('/')
def main():
    return render_template('main.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        if username in users and users[username]['password'] == password:
            user = User(username)
            login_user(user)
            flash('Login successful!', 'success')
            return redirect(url_for('form'))
        else:
            flash('Invalid credentials, please try again.', 'error')
    
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('You have been logged out.', 'success')
    return redirect(url_for('login'))

# Constants
COVER_LETTER_TEMPLATE = "Cover_letterr.docx"
FINAL_FILE_DOCX_FILENAME = "Final_Cover_letter_with_table_{}.docx"
FINAL_FILE_PDF_FILENAME = "Final_Cover_letter_with_table_{}.pdf"
PDFS_DIRECTORY = os.path.join(app.root_path, 'pdfs')

# Ensure the PDFs directory exists
if not os.path.exists(PDFS_DIRECTORY):
    os.makedirs(PDFS_DIRECTORY)

def generate_cover_letter(context):
    try:
        today_date = datetime.datetime.today().strftime('%B %d, %Y')
        context['today_date'] = today_date

        doc = DocxTemplate(COVER_LETTER_TEMPLATE)
        doc.render(context)
        
        temp_filename = "Temp_Cover_letter.docx"
        doc.save(temp_filename)

        return temp_filename
    except Exception as e:
        logging.error("Error in generate_cover_letter: %s", e)
        raise

def create_and_insert_table(doc, target_index, records):
    try:
        num_cols = 5
        table = doc.add_table(rows=len(records) + 1, cols=num_cols)

        headers = ["S.no", "Description", "Rate", "Quantity", "Amount"]
        for i, header_text in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header_text
            shading_color = "808080"
            cell._element.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_color}" w:val="clear"/>'))
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].runs[0].font.size = Pt(12)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row in table.rows:
            row.cells[0].width = Inches(0.3)
            row.cells[1].width = Inches(3.0)
            row.cells[2].width = Inches(0.5)
            row.cells[3].width = Inches(0.5)
            row.cells[4].width = Inches(0.5)

        for i, record in enumerate(records, start=1):
            for j, header_text in enumerate(headers):
                cell = table.cell(i, j)
                if header_text == "Amount":
                    rate = float(record[2])
                    quantity = float(record[3])
                    amount = rate * quantity
                    cell.text = str(amount)
                else:
                    cell.text = str(record[j])
                
                shading_color = "D3D3D3" if i % 2 == 0 else "FFFFFF"
                cell._element.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_color}" w:val="clear"/>'))
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                cell.paragraphs[0].runs[0].font.size = Pt(12)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.style = 'Table Grid'
        paragraph = doc.paragraphs[target_index]
        paragraph.insert_paragraph_before()._p.addnext(table._tbl)
    except Exception as e:
        logging.error("Error in create_and_insert_table: %s", e)
        raise

def convert_docx_to_html(file_path):
    try:
        with open(file_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value
            return html
    except Exception as e:
        logging.error("Error in convert_docx_to_html: %s", e)
        raise

@app.route('/form', methods=['GET', 'POST'])
@login_required
def form():
    try:
        if request.method == 'POST':
            context = {
                'offer_name': request.form['offer_name'],
                'offer_no': request.form['offer_no'],
                'company_name': request.form['company_name'],
                'city_name': request.form['city_name'],
                'state_name': request.form['state_name'],
                'manager_name': request.form['manager_name'],
                'enquiry_sub': request.form['enquiry_sub'],
                'delivery_dates': request.form['delivery_dates'],
                'your_name': request.form['your_name'],
                'contact_no': request.form['contact_no']
            }
            records_count = int(request.form['records_count'])
            records = []
            for i in range(records_count):
                sn = request.form[f'sn_{i}']
                description = request.form[f'description_{i}']
                rate = float(request.form[f'rate_{i}'])
                quantity = float(request.form[f'quantity_{i}'])
                records.append((sn, description, rate, quantity))

            cover_letter_file = generate_cover_letter(context)
            doc = Document(cover_letter_file)

            target_text = "Annexure II-Commercial Terms and Conditions."
            target_index = None
            for i, paragraph in enumerate(doc.paragraphs):
                if target_text in paragraph.text:
                    target_index = i
                    break

            if target_index is not None:
                create_and_insert_table(doc, target_index, records)
                unique_suffix = datetime.datetime.now().strftime('%Y%m%d%H%M%S')
                final_docx_file = FINAL_FILE_DOCX_FILENAME.format(unique_suffix)
                doc.save(final_docx_file)

                if not os.path.exists(final_docx_file):
                    logging.error("Error: Final DOCX file not created.")
                    return "Error: Final DOCX file not created."

                pythoncom.CoInitialize()
                try:
                    final_pdf_file = os.path.join(PDFS_DIRECTORY, FINAL_FILE_PDF_FILENAME.format(unique_suffix))
                    convert(final_docx_file, final_pdf_file)
                finally:
                    pythoncom.CoUninitialize()

                if not os.path.exists(final_pdf_file):
                    logging.error("Error: Final PDF file not created.")
                    return "Error: Final PDF file not created."

                html_content = convert_docx_to_html(final_docx_file)

                pdfs_metadata.insert_one({
                    'filename': FINAL_FILE_PDF_FILENAME.format(unique_suffix),
                    'user_id': current_user.id,
                    'created_at': datetime.datetime.now()
                })

                return render_template('preview.html', html_content=html_content, filename=final_pdf_file)
            else:
                return "Error: Target paragraph not found in the document."
    except Exception as e:
        logging.error("Error in form route: %s", e)
        return str(e)

    return render_template('form.html')

@app.route('/list_pdfs')
@login_required
def list_pdfs():
    try:
        if current_user.id == 'admin':
            pdf_files = list(pdfs_metadata.find({}, {'_id': 0, 'filename': 1}))
        else:
            pdf_files = list(pdfs_metadata.find({'user_id': current_user.id}, {'_id': 0, 'filename': 1}))

        pdf_files = [pdf['filename'] for pdf in pdf_files]
        return render_template('list_pdfs.html', pdf_files=pdf_files)
    except Exception as e:
        logging.error("Error in list_pdfs route: %s", e)
        return str(e)

@app.route('/view/<filename>')
@login_required
def view_pdf(filename):
    try:
        file_metadata = pdfs_metadata.find_one({'filename': filename})
        if not file_metadata:
            return "Error: File metadata not found."
        
        user = db.users.find_one({"username": current_user.id})
        is_admin = user.get('is_admin', False) if user else False
        return render_template('view_pdfs.html', filename=filename, username=current_user.id if is_admin else None)
    except Exception as e:
        logging.error("Error in view_pdf route: %s", e)
        return str(e)

@app.route('/download_pdf/<filename>')
@login_required
def download_pdf(filename):
    try:
        file_path = os.path.join(PDFS_DIRECTORY, filename)
        if not os.path.exists(file_path):
            return "Error: PDF file not found."

        return send_file(file_path, as_attachment=True)
    except Exception as e:
        logging.error("Error in download_pdf route: %s", e)
        return str(e)
    
@app.route('/serve_pdf/<filename>')
@login_required
def serve_pdf(filename):
    try:
        file_path = os.path.join(PDFS_DIRECTORY, filename)
        if not os.path.exists(file_path):
            return "Error: PDF file not found."

        return send_file(file_path, as_attachment=False)
    except Exception as e:
        logging.error("Error in serve_pdf route: %s", e)
        return str(e)

if __name__ == '__main__':
    app.run("0.0.0.0")

